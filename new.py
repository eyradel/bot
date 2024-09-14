import streamlit as st
import os
import fitz  # PyMuPDF for PDF files
import docx  # For handling .docx files
from transformers import AutoTokenizer, AutoModelForQuestionAnswering, pipeline
from langchain_community.llms import HuggingFacePipeline
from dotenv import load_dotenv
from langchain.chains import RetrievalQA
from langchain.vectorstores import FAISS
from langchain.embeddings import HuggingFaceEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.docstore.document import Document

load_dotenv()

st.set_page_config(layout="wide")

# Initialize Hugging Face embeddings model
embedding_model = HuggingFaceEmbeddings()

# FAISS vector store initialization using from_documents
def init_faiss(documents):
    return FAISS.from_documents(documents, embedding_model)

# Text splitter for long documents
text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=150)

# Function to extract text from multiple PDFs or Word files
def extract_text_from_files(uploaded_files):
    extracted_texts = []
    for uploaded_file in uploaded_files:
        if uploaded_file.name.endswith(".pdf"):
            # Handle PDF files
            doc = fitz.open(stream=uploaded_file.read(), filetype="pdf") 
            text = ""
            for page in doc:
                text += page.get_text()  # Extract text from each page
            extracted_texts.append(text)
        elif uploaded_file.name.endswith(".docx"):
            # Handle Word files
            doc = docx.Document(uploaded_file)
            text = "\n".join([para.text for para in doc.paragraphs])
            extracted_texts.append(text)
    return " ".join(extracted_texts)

# Function to split text and add to the vector store
def split_and_store_text(text):
    chunks = text_splitter.split_text(text)
    documents = [Document(page_content=chunk) for chunk in chunks]
    return documents

# Hugging Face Question-Answering model setup
tokenizer = AutoTokenizer.from_pretrained("distilbert-base-uncased-distilled-squad")
qa_model = AutoModelForQuestionAnswering.from_pretrained("distilbert-base-uncased-distilled-squad")
question_answerer = pipeline("question-answering", model=qa_model, tokenizer=tokenizer)

# LLM pipeline using Hugging Face
local_llm = HuggingFacePipeline(pipeline=question_answerer)

# Setup RetrievalQA Chain with FAISS retriever
def init_retrieval_chain(documents):
    faiss_index = init_faiss(documents)
    retriever = faiss_index.as_retriever()
    return RetrievalQA.from_chain_type(llm=local_llm, chain_type="stuff", retriever=retriever)

@st.cache_resource
def llm_pipeline():
    return local_llm

def process_summary(extracted_text):
    documents = split_and_store_text(extracted_text)  # Split the text into chunks
    retrieval_chain = init_retrieval_chain(documents)  # Initialize chain with documents
    
    # Get the context from the vector store
    context = " ".join([doc.page_content for doc in documents])
    
    # Prepare the question for summarization
    instruction = "Summarize the following text."
    
    # Pass both the question and context to the pipeline
    result = question_answerer(question=instruction, context=context)
    return result['answer']

def process_template(extracted_text):
    first_visit_date = "placeholder for first visit date"
    last_visit_date = "placeholder for last visit date"
    documents = split_and_store_text(extracted_text)  # Store text in the vector store
    retrieval_chain = init_retrieval_chain(documents)  # Initialize chain with documents
    
    # Get the context from the vector store
    context = " ".join([doc.page_content for doc in documents])
    
    # Prepare the question for template generation
    instruction = f"The first visit date was {first_visit_date} and the last visit date was {last_visit_date}. Analyze the text."
    
    # Pass both the question and context to the pipeline
    result = question_answerer(question=instruction, context=context)
    return result['answer']

def main():
    st.title("Document Analyzer with FAISS-based RAG")

    # Upload multiple PDF or Word files
    uploaded_files = st.file_uploader("Upload PDF or Word files", type=["pdf", "docx"], accept_multiple_files=True)

    if uploaded_files:
        with st.spinner("Extracting Text..."):
            extracted_text = extract_text_from_files(uploaded_files)

        st.success("Text Extracted")

        col1, col2 = st.columns([1, 1])

        # Display the extracted text
        with col1:
            st.subheader("Extracted Text")
            st.write(extracted_text)  # Display the entire extracted text
            with st.spinner("Summarizing..."):
                summary = process_summary(extracted_text)
            st.success("Summarized Text")
            st.write(summary)  # Display the summarized text

        # Display the generated template
        with col2:
            with st.spinner("Generating Template"):
                template = process_template(extracted_text)
            st.success("Generated Template")
            st.write(template)

    st.write("Upload your files")

if __name__ == "__main__":
    main()
